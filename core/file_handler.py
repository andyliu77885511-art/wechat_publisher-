"""
file_handler.py — 文件上传处理 + ffmpeg 音频提取 + 文档文本提取
"""
import os
import subprocess
import uuid
from pathlib import Path

import config

UPLOAD_DIR = config.UPLOAD_DIR


def save_uploaded_file(uploaded_file) -> tuple[str, str]:
    """
    保存 Streamlit 上传的文件到本地
    返回 (file_path, file_type)
    """
    Path(UPLOAD_DIR).mkdir(parents=True, exist_ok=True)
    suffix = Path(uploaded_file.name).suffix.lower()
    uid = str(uuid.uuid4())
    dest = os.path.join(UPLOAD_DIR, f"{uid}{suffix}")
    with open(dest, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return dest, suffix.lstrip(".")


def extract_audio(video_path: str) -> str:
    """
    用 ffmpeg 从视频文件提取音频，保存为 mp3
    返回音频文件路径
    """
    audio_path = video_path.rsplit(".", 1)[0] + "_audio.mp3"
    cmd = [
        "ffmpeg", "-y",
        "-i", video_path,
        "-vn",
        "-acodec", "libmp3lame",
        "-ar", "16000",
        "-ac", "1",
        "-ab", "64k",
        audio_path
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"ffmpeg 提取音频失败：{result.stderr[-500:]}")
    return audio_path


def get_audio_path(file_path: str, file_type: str) -> str:
    """
    如果是视频文件则提取音频，如果已是音频直接返回
    """
    video_types = {"mp4", "avi", "mov", "mkv", "flv", "wmv"}
    if file_type in video_types:
        return extract_audio(file_path)
    return file_path


def get_duration_seconds(audio_path: str) -> int:
    """
    获取音频时长（秒）
    """
    try:
        cmd = [
            "ffprobe", "-v", "error",
            "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1",
            audio_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        return int(float(result.stdout.strip()))
    except Exception:
        return 0


def estimate_transcribe_minutes(duration_seconds: int) -> str:
    """
    根据音频时长估算转录耗时
    """
    mins = duration_seconds / 60
    if mins <= 10:
        return "约1分钟"
    elif mins <= 30:
        return "约2-4分钟"
    elif mins <= 60:
        return "约4-8分钟"
    else:
        return f"约{int(mins / 8)}-{int(mins / 6)}分钟"


def extract_text_from_doc(file_path: str, file_type: str) -> str:
    """
    从文档文件（pdf/txt/docx/md）中提取纯文本
    返回提取的文本内容
    提取失败抛出 RuntimeError
    """
    if file_type == "txt":
        return _extract_text_from_txt(file_path)
    elif file_type == "md":
        return _extract_text_from_md(file_path)
    elif file_type == "pdf":
        return _extract_text_from_pdf(file_path)
    elif file_type == "docx":
        return _extract_text_from_docx(file_path)
    else:
        raise RuntimeError(f"不支持的文档格式：{file_type}")


def _extract_text_from_txt(file_path: str) -> str:
    """读取纯文本文件"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        # 尝试其他编码
        for enc in ["gbk", "gb2312", "latin-1"]:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    return f.read()
            except Exception:
                continue
        raise RuntimeError(f"无法解码文本文件（尝试了 utf-8/gbk/gb2312）")
    except Exception as e:
        raise RuntimeError(f"读取文本文件失败：{e}")


def _extract_text_from_md(file_path: str) -> str:
    """Markdown 本质上就是纯文本，直接读取"""
    return _extract_text_from_txt(file_path)


def _extract_text_from_pdf(file_path: str) -> str:
    """使用 pypdf 提取 PDF 文本"""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise RuntimeError("缺少依赖 pypdf，请先安装：pip install pypdf")

    try:
        reader = PdfReader(file_path)
        texts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                texts.append(page_text)
        if not texts:
            raise RuntimeError("PDF 中未提取到文字内容，可能是扫描版或图片型 PDF")
        return "\n".join(texts)
    except Exception as e:
        if "pypdf" in str(e) or "PdfReader" in str(e):
            raise RuntimeError(f"PDF 解析失败：{e}")
        raise RuntimeError(f"PDF 提取失败：{e}")


def _extract_text_from_docx(file_path: str) -> str:
    """使用 python-docx 提取 Word 文档文本"""
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("缺少依赖 python-docx，请先安装：pip install python-docx")

    try:
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        # 尝试读取表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        if not paragraphs:
            raise RuntimeError("Word 文档中未提取到文字内容")
        return "\n".join(paragraphs)
    except Exception as e:
        raise RuntimeError(f"Word 文档提取失败：{e}")
